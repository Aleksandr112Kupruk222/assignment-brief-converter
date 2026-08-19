import io

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


client = TestClient(app)


def docx_bytes(label="Assignment Title", value="A title"):
    stream = io.BytesIO(); doc = Document(); table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = label; table.cell(0, 1).text = value; doc.save(stream); return stream.getvalue()


def test_docx_upload_and_generation():
    response = client.post("/api/analyse", files={
        "old_document": ("old.docx", docx_bytes(value="Original title"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        "new_document": ("new.docx", docx_bytes(value="Example title"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    assert response.status_code == 200
    data = response.json(); mapping = data["mappings"][0]
    generated = client.post("/api/generate", data={"session_id": data["session_id"], "mappings": __import__('json').dumps({mapping["target_id"]: mapping["source_id"]}), "manual_values": "{}"})
    assert generated.status_code == 200
    assert generated.content.startswith(b"PK")


def test_unsupported_corrupt_empty_and_duplicate_fail_gracefully():
    unsupported = client.post("/api/analyse", files={"old_document": ("old.pdf", b"x"), "new_document": ("new.docx", docx_bytes())})
    assert unsupported.status_code == 415
    corrupt = client.post("/api/analyse", files={"old_document": ("old.docx", b"not docx"), "new_document": ("new.docx", docx_bytes())})
    assert corrupt.status_code == 422
    same = docx_bytes()
    duplicate = client.post("/api/analyse", files={"old_document": ("old.docx", same), "new_document": ("new.docx", same)})
    assert duplicate.status_code == 400
    empty_stream = io.BytesIO(); Document().save(empty_stream)
    empty = client.post("/api/analyse", files={"old_document": ("old.docx", empty_stream.getvalue()), "new_document": ("new.docx", docx_bytes())})
    assert empty.status_code == 422

