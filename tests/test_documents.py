from pathlib import Path

from docx import Document

from app.documents import extract_fields, generate_document, match_fields


def make_old(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Assignment Title: Design a Campaign")
    doc.add_heading("Vocational Scenario", level=1)
    doc.add_paragraph("You are working for a local design studio.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Unit Title"
    table.cell(0, 1).text = "Digital Design"
    table.cell(1, 0).text = "Submission Deadline"
    table.cell(1, 1).text = "30 June 2027"
    doc.save(path)


def make_new(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Assignment Title"
    table.cell(0, 1).text = "Example title"
    table.cell(1, 0).text = "Unit / Module"
    table.cell(1, 1).text = "Example unit"
    table.cell(2, 0).text = "Submission Date"
    table.cell(2, 1).text = "Example date"
    doc.add_heading("Vocational Scenario", level=1)
    doc.add_paragraph("Example scenario text")
    doc.save(path)


def test_paragraph_and_table_extraction(tmp_path):
    path = tmp_path / "old.docx"; make_old(path)
    fields = extract_fields(Document(path), "old")
    assert any(f.canonical == "assignment_title" and f.value == "Design a Campaign" for f in fields)
    assert any(f.canonical == "unit_title" and f.value == "Digital Design" for f in fields)


def test_section_detection(tmp_path):
    path = tmp_path / "old.docx"; make_old(path)
    fields = extract_fields(Document(path), "old")
    assert any(f.canonical == "scenario" and "local design studio" in f.value for f in fields)


def test_mapping_and_missing_fields(tmp_path):
    old, new = tmp_path / "old.docx", tmp_path / "new.docx"; make_old(old); make_new(new)
    old_fields = extract_fields(Document(old), "old")
    new_fields = extract_fields(Document(new), "new", include_slots=True)
    mappings = match_fields(old_fields, new_fields)
    title = next(m for m in mappings if m["target_label"] == "Assignment Title")
    assert title["source_id"] is not None
    old_fields = [f for f in old_fields if f.canonical != "submission_date"]
    mappings = match_fields(old_fields, new_fields)
    date = next(m for m in mappings if m["target_label"] == "Submission Date")
    assert date["source_id"] is None


def test_generation_preserves_template_and_replaces_values(tmp_path):
    old, new, output = tmp_path / "old.docx", tmp_path / "new.docx", tmp_path / "out.docx"
    make_old(old); make_new(new)
    old_fields = extract_fields(Document(old), "old")
    new_fields = extract_fields(Document(new), "new", include_slots=True)
    mappings = match_fields(old_fields, new_fields)
    selections = {m["target_id"]: m["source_id"] for m in mappings}
    generate_document(new, output, old_fields, new_fields, selections, {})
    generated = Document(output)
    assert generated.tables[0].cell(0, 1).text == "Design a Campaign"
    assert generated.tables[0].cell(1, 1).text == "Digital Design"
    assert generated.tables[0].cell(2, 1).text == "30 June 2027"
    assert generated.paragraphs[-1].text == "You are working for a local design studio."

