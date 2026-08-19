from __future__ import annotations

import re
from dataclasses import dataclass, asdict
from difflib import SequenceMatcher
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


FIELD_ALIASES = {
    "assignment_title": ("assignment title", "title of assignment", "assignment name"),
    "qualification": ("qualification", "programme", "course"),
    "unit_title": ("unit title", "module title", "unit / module", "unit/module"),
    "unit_number": ("unit number", "unit no", "module number", "module code", "unit code"),
    "assessor": ("assessor", "teacher", "tutor"),
    "student": ("student", "learner", "student name", "learner name"),
    "issue_date": ("issue date", "date issued", "start date"),
    "submission_date": ("submission date", "deadline", "due date", "hand in date"),
    "scenario": ("assignment scenario", "vocational scenario", "scenario", "context"),
    "purpose": ("assignment purpose", "purpose", "aim of assignment"),
    "learning_outcomes": ("learning outcomes", "learning outcome", "learning aims", "learning aim"),
    "assessment_criteria": ("assessment criteria", "grading criteria", "criteria covered", "criteria"),
    "tasks": ("assignment tasks", "assignment activity", "activities", "tasks", "task"),
    "evidence": ("evidence required", "evidence", "submission evidence"),
    "resources": ("resources", "references", "recommended resources"),
    "internal_verifier": ("internal verifier", "iv name", "verified by"),
    "programme_code": ("programme code", "course code"),
}


def normalise(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()


def canonical_for(label: str) -> str | None:
    value = normalise(label)
    if not value:
        return None
    best: tuple[float, str] = (0.0, "")
    for key, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            alias_n = normalise(alias)
            score = SequenceMatcher(None, value, alias_n).ratio()
            if alias_n in value or value in alias_n:
                score = max(score, 0.9)
            if score > best[0]:
                best = (score, key)
    return best[1] if best[0] >= 0.72 else None


@dataclass
class Field:
    id: str
    label: str
    canonical: str | None
    value: str
    source: str
    locator: dict | None = None

    def public(self) -> dict:
        return asdict(self)


def _text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def _is_heading(paragraph: Paragraph) -> bool:
    style = paragraph.style.name.lower() if paragraph.style else ""
    return style.startswith("heading") or style in {"title", "subtitle"}


def _add_field(fields: list[Field], label: str, value: str, source: str, locator: dict | None = None) -> None:
    value = value.strip()
    label = label.strip()
    if not label or not value:
        return
    fields.append(Field(f"{source}-{len(fields)}", label, canonical_for(label), value, source, locator))


def extract_fields(document: DocumentType, source: str, include_slots: bool = False) -> list[Field]:
    fields: list[Field] = []
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        text = _text(paragraph)
        if not text:
            continue
        if ":" in text:
            label, value = text.split(":", 1)
            if canonical_for(label) and value.strip():
                _add_field(fields, label, value, source, {"kind": "paragraph", "index": index} if include_slots else None)
                continue
        if _is_heading(paragraph):
            following: list[str] = []
            first_content_index: int | None = None
            for next_index, next_paragraph in enumerate(paragraphs[index + 1 :], start=index + 1):
                if _is_heading(next_paragraph) and _text(next_paragraph):
                    break
                if _text(next_paragraph):
                    if first_content_index is None:
                        first_content_index = next_index
                    following.append(_text(next_paragraph))
            if following and first_content_index is not None:
                _add_field(fields, text, "\n\n".join(following), source,
                           {"kind": "paragraph", "index": first_content_index} if include_slots else None)

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) < 2:
                continue
            label = cells[0].text.strip()
            value = "\n".join(c.text.strip() for c in cells[1:] if c.text.strip())
            if label and (canonical_for(label) or include_slots):
                _add_field(fields, label, value, source,
                           {"kind": "cell", "table": table_index, "row": row_index, "cell": 1} if include_slots else None)
    return deduplicate(fields)


def deduplicate(fields: Iterable[Field]) -> list[Field]:
    result: list[Field] = []
    seen: set[tuple[str, str]] = set()
    for field in fields:
        key = (field.canonical or normalise(field.label), normalise(field.value))
        if key not in seen:
            seen.add(key)
            result.append(field)
    return result


def inspect_docx(path: Path, source: str, include_slots: bool = False) -> tuple[DocumentType, list[Field]]:
    try:
        document = Document(path)
    except Exception as exc:
        raise ValueError("The file is not a readable Word .docx document.") from exc
    visible = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    visible.extend(c.text.strip() for t in document.tables for r in t.rows for c in r.cells if c.text.strip())
    if not visible:
        raise ValueError("The Word document is empty.")
    return document, extract_fields(document, source, include_slots)


def match_fields(old_fields: list[Field], new_fields: list[Field]) -> list[dict]:
    mappings: list[dict] = []
    for target in new_fields:
        best: tuple[float, Field | None] = (0.0, None)
        for source in old_fields:
            score = SequenceMatcher(None, normalise(source.label), normalise(target.label)).ratio()
            if source.canonical and source.canonical == target.canonical:
                score = 1.0
            if score > best[0]:
                best = (score, source)
        source = best[1] if best[0] >= 0.52 else None
        mappings.append({
            "target_id": target.id,
            "target_label": target.label,
            "source_id": source.id if source else None,
            "source_label": source.label if source else None,
            "preview": source.value[:220] if source else "",
            "confidence": round(best[0], 2) if source else 0,
        })
    return mappings


def _replace_paragraph(paragraph: Paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _replace_cell(cell: _Cell, value: str) -> None:
    paragraphs = cell.paragraphs
    _replace_paragraph(paragraphs[0], value)
    for paragraph in paragraphs[1:]:
        _replace_paragraph(paragraph, "")


def generate_document(template_path: Path, output_path: Path, old_fields: list[Field], new_fields: list[Field],
                      selections: dict[str, str | None], manual_values: dict[str, str]) -> None:
    document = Document(template_path)
    old_by_id = {field.id: field for field in old_fields}
    new_by_id = {field.id: field for field in new_fields}
    for target_id, target in new_by_id.items():
        source_id = selections.get(target_id)
        value = manual_values.get(target_id, "").strip()
        if not value and source_id in old_by_id:
            value = old_by_id[source_id].value
        if not value or not target.locator:
            continue
        locator = target.locator
        if locator["kind"] == "paragraph" and locator["index"] < len(document.paragraphs):
            _replace_paragraph(document.paragraphs[locator["index"]], value)
        elif locator["kind"] == "cell":
            table: Table = document.tables[locator["table"]]
            _replace_cell(table.rows[locator["row"]].cells[locator["cell"]], value)
    document.save(output_path)

